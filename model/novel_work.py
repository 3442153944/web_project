import os
import re

import tornado.ioloop
import tornado.web
import json
from model.connect_sqlsever import *
from tornado.httpclient import AsyncHTTPClient
from model.CORSMixin import *
from datetime import datetime
from docx import Document
from docx.shared import Pt


class get_novel_work(tornado.web.RequestHandler, CORSMixin):
    conn = connMysql()

    def post(self):
        try:
            self.set_status(200)
            self.set_header('Content-Type', 'application/json')
            data = json.loads(self.request.body.decode('utf-8'))
            work_id = data["work_id"]
            sql = "select * from novel_work where work_id=%s"
            conn = self.conn.connect()
            cursor = conn.cursor()
            cursor.execute(sql, (work_id,))
            result = cursor.fetchall()
            if result:
                column_names=[desc[0] for desc in cursor.description]
                result_list=[dict(zip(column_names,row)) for row in result]
                self.write(json.dumps({"status":"success", "data": result_list}))
            else:
                self.write(json.dumps({"status":"error","msg":"未找到该作品"}))
            conn.close()
        except Exception as e:
            self.write(json.dumps({"status":"error","msg":"服务器错误"}))
            print(e)

    def get(self):
        self.set_status(200)
        work_id = self.get_argument('work_id')
        work_title = self.get_argument('title_text')
        work_name = self.get_argument('work_name')
        content = self.get_word_content(work_name, work_title)
        self.write(json.dumps({"work_content": content}))

    def get_word_content(self, name, work_title):
        file_src = 'H:/web_preject/novel_work/' + name + "/" + work_title + ".docx"
        doc = Document(file_src)
        content = ""
        for para in doc.paragraphs:
            content += para.text + "<br>"
            # print(para.text)
        return content

    def is_title(self, paragraph):
        if paragraph.text.startswith("第") and "章" in paragraph.text:
            return True
        else:
            return False

    def get_title_name(self, title_text):
        # 假设标题格式为"第......章 标题名称"
        return title_text.split(" ", 1)[1]

    def get_work_list(self, work_name):
        work_file = 'H:/web_preject/novel_work/' + work_name + '.docx'
        title_list = []

        try:
            # 打开 Word 文档
            doc = Document(work_file)
            count = 0
            temp_text = ''

            # 提取标题
            for para in doc.paragraphs:
                temp_text = para.text
                # 判断是否为标题
                if "第" in temp_text and "章" in temp_text:
                    title_list.append(temp_text.strip())
                    temp_text = ''
                    count += 1

            # 如果没有找到标题，则通知
            if not title_list:
                print("未找到任何标题")
            # print(count)
            # print("标题列表:", title_list)
            return title_list

        except FileNotFoundError:
            print("文件未找到:", work_file)
            return None
        except Exception as e:
            print("发生错误:", e)
            return None

    def chose_title(self, title, work_name):
        work_file = 'H:/web_preject/novel_work/' + work_name + '.docx'
        title_content = ''
        return title_content

    import os

    def get_file_name_list(self, name):
        work_file = "H:/web_preject/novel_work/" + name
        file_list = []
        try:
            for file_name in os.listdir(work_file):
                if os.path.isfile(os.path.join(work_file, file_name)):
                    # 分割文件名和后缀名
                    file_name_no_extension, _ = os.path.splitext(file_name)
                    file_list.append(file_name_no_extension)  # 将文件名添加到列表中

        except Exception as e:
            print(e)

        # 对文件名进行排序
        file_list.sort(key=self.sort_by_number)
        return file_list

    def sort_by_number(self, filename):
        # 从文件名中提取数字部分
        num_str = filename.split('-')[-1].split('.')[0]
        # 转换为整数
        num = int(num_str)
        # 返回提取的数字
        return num


class get_novel_content(tornado.web.RequestHandler, CORSMixin):
    conn = connMysql()
    work_name = ''
    work_content = ''

    def post(self):
        self.set_status(200)
        self.set_header('Content-Type', 'application/json')
        data = json.loads(self.request.body.decode('utf-8'))

        try:
            print(data)
            print(data["work_name"])
            title = data["work_title"]
            print(title)
            self.work_name = data["work_name"]
            # self.get_work_content(title)  # 调用获取工作内容的方法，并传递标题作为参数
            # print(self.work_content + 'test')
            self.write(json.dumps({"work_content": 'test'}))
        except Exception as e:
            print(e)
            print('发送正文内容失败')

    def get_work_content(self, title):
        work_file = 'H:/web_preject/novel_work/' + self.work_name + '.docx'

        # 清理并标准化传入的标题
        title = re.sub(r'.*?(第[\u4e00-\u9fa5]+章).*', r'\1', title)
        title = title.strip()

        try:
            if not title:  # 如果标题为空，则默认返回第一章的内容
                title = "第一章"

            doc = Document(work_file)

            self.work_content = ""  # 清空正文内容

            collecting_content = False  # 标记是否开始收集正文内容

            for i, para in enumerate(doc.paragraphs):
                temp_text = para.text
                print(temp_text)
                # 检查当前段落是否为章节标题
                if temp_text.startswith("第") and "章" in temp_text:
                    chapter_title = re.sub(r'.*?(第[\u4e00-\u9fa5]+章).*', r'\1', temp_text)
                    chapter_title = chapter_title.strip()

                    # 如果找到了匹配的章节标题，则开始收集正文内容
                    if chapter_title == title:
                        collecting_content = True
                    else:
                        # 如果不是匹配的章节标题，则停止收集正文内容
                        collecting_content = False

                # 如果正在收集正文内容，则将当前段落添加到内容中
                if collecting_content:
                    self.work_content += para.text + "<br>"

                # 检查是否已经到达了列表的末尾
                if i < len(doc.paragraphs) - 1:
                    next_para = doc.paragraphs[i + 1]
                    if next_para.text.startswith("第") and "章" in next_para.text:
                        break  # 如果是新章节的开始，则停止收集

        except Exception as e:
            print(e)

        return self.work_content



